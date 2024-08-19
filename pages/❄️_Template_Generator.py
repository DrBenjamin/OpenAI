##### `❄️_Template_Generator.py`
##### BAS Anzeigen Generator
##### Please reach out to benjamin.gross1@adesso.de for any questions
#### Loading needed Python libraries
import streamlit as st
import pandas as pd
import sys
import datetime
import requests
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from snowflake.snowpark import Session
from langchain_community.chat_message_histories import StreamlitChatMessageHistory
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_openai import ChatOpenAI
    
# Establish Snowflake session
@st.cache_resource
def create_session():
    return Session.builder.configs(st.secrets.snowflake).create()

session = create_session()
st.success("Datenbankverbindung erfolgreich hergestellt.")

# Write data table
def write_data(data, table_name, database, schema):
    # Write data to table
    session.write_pandas(data, table_name=table_name, database=database, schema=schema, overwrite=True)
    st.success("Daten erfolgreich geschrieben.")

# Load data table
@st.cache_data
def load_data(table_name):
    # Read in data table
    table = session.table(table_name)

    # Collect the results. This will run the query and download the data
    table = table.collect()
    st.success("Daten erfolgreich gelesen.")
    return pd.DataFrame(table)

# Web Scraper
def web_scraper(url):
    info_page = requests.get(url)
    info_soup = BeautifulSoup(info_page.content, 'html.parser')
    info = info_soup.get_text()
    info = info.replace('\n', ' ')
    return info

# Word export
def export_doc(data):
    document = Document()

    # Adding Image
    document.add_picture('images/gwq_logo_header.png')
    
    # Adding centered text
    centered_text = f"""\n\n\n
                        {cloud} 
                        inkl.
                        {service_1}
                        {service_2}
                        \n\n
                        Anforderungen 
                        Regulatorik
                        """
    paragraph = document.add_paragraph(centered_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(24)
    centered_date = f"""\n\n
                        {datetime.date.today()}"""
    paragraph = document.add_paragraph(centered_date)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)
    document.add_page_break()
    
    # Adding table of changes
    def set_table_borders(table):
        tbl = table._element
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Border color
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)
    paragraph = document.add_paragraph('Änderungshistorie des Dokumentes')
    table = document.add_table(rows=5, cols=7)
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    headers = ['Nr.', 'Datum', 'Version', 'Kapitel', 'Beschreibung der Änderung', 'Autor', 'Bearbeitungszustand']
    for i, header in enumerate(headers):
        run = hdr_cells[i].add_paragraph().add_run(header)
        run.bold = True
        run.font.size = Pt(10)
    row_cells = table.rows[1].cells
    contents = [str(1), str(datetime.date.today()), '1.0', 'alle', 'Erstellung des Dokuments', 'Groß, Benjamin', 'Erl.']
    for i, content in enumerate(contents):
        run = row_cells[i].add_paragraph().add_run(content)
        run.font.size = Pt(10)
    document.add_page_break()

    # Adding Table of Contents
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run('Inhaltsverzeichnis')
    run.font.size = Pt(16)

    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = " (Rechts-click um Inhaltsverzeichnis hinzuzufügen - Update Feld)"

    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    document.add_page_break()

    # Writing paragraphs
    for index, row in data.iterrows():
        document.add_heading(f"{row['PARAGRAPH']} - {row['PARAGRAPH_TITLE']}", level=len(row['PARAGRAPH'].replace('.', '')))
        paragraph = document.add_paragraph()
        paragraph.add_run(f"\n{row['PARAGRAPH_TEXT']}\n")
        paragraph.style.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Download button
    buffer = BytesIO()
    document.save(buffer)
    st.toast('Das Dokument ist fertig!', icon ='📃')
    st.download_button(label='Download Template', data=buffer, file_name='BAS_Anzeige_Template.docx', mime='application/vnd.openxmlformats')

# Title
st.title('❄️ Template Generator')
st.write(f"Streamlit Version: {st.__version__}")
st.write(f"Python Version: {sys.version}")

# Display data table
with st.expander("Datenbankinhalt"):
    df = load_data('OPENAI_DATABASE.PUBLIC.ANZEIGE_PRE')
    st.dataframe(df)
    paragraphs = load_data('OPENAI_DATABASE.PUBLIC.ANZEIGE_PARAGRAPHS')
    st.dataframe(paragraphs)
    options = load_data('OPENAI_DATABASE.PUBLIC.ANZEIGE_OPTIONS')
    st.dataframe(options)

# Sidebar
sidebar = st.sidebar
with sidebar:
    kunde = st.text_input("Kunde:", value="GWQ ServicePlus AG")
    web = st.toggle("Webscraper", True)
    kunde_url = st.text_input("Kunden-Webseite (z.B. `Über uns`):", value="https://www.gwq-serviceplus.de/ueber-uns")
    if web:
        kunde_info = web_scraper(kunde_url)
    cloud = st.selectbox("Cloud:", ["AWS", "Azure", "Google Cloud"], index=2)
    service_1 = 'Google Cloud Vision'
    service_2 = 'Google Translate'
    on = st.toggle("OpenAI ChatGPT", True)
    system = st.text_input("System:", value = f"Du erstellst einzelne Absätze einer Anzeige beim Bundesamt für Soziale Sicherung über die Verarbeitung von Sozialdaten im Auftrag (AVV) nach § 80 Zehntes Sozialgesetzbuch (SGB X). Tausche die Platzhalter (z.B. <Kunde>) durch die entsprechenden Inhalte aus und gebe nur den verbesserten Text in einer sachlichen und formellen Form aus und verzichte auf Phrasen wie z.B. 'Vielen Dank für die Informationen. Hier sind die angepassten Absätze für die Anzeige beim Bundesamt für Soziale Sicherung:'.")
    if not on:
      st.markdown("Local Server Configuration")
      url = st.text_input("URL:", value="http://localhost")
      port = st.number_input("Port:", value=1234, min_value=1, max_value=65535)

with st.form("form"):
    st.title("Konfiguration")
    st.write("Bitte fülle die folgenden Felder aus, um bestmöglich ein Template generieren zu können.")
    base_information = st.container(border=True)
    base_information.header("Grundsätzliche Informationen")
    with base_information.container(border=True):
        st.subheader("Anbieterinformationen")
        base_information_0_0 = st.text_input("Name des Auftragsverarbeiters?", key='option_0_0', value=str(options['OPTION_TEXT'].iloc[0]))
    options_1 = st.container(border=True)
    options_1.header("Sozialdaten")
    with options_1.container(border=True):
        st.subheader("Details zu den Sozialdaten")
        options_1_0a = st.toggle("Sollen Sozialdaten verarbeitet werden?", key='option_1_0a', value=eval(options['OPTION_TEXT'].iloc[1]))
        options_1_0b = st.toggle("Sollen diese persistent gespeichert werden?", key='option_1_0b', value=eval(options['OPTION_TEXT'].iloc[2]))
        options_1_1 = st.text_input("Welche Art von Sozialdaten?", key='option_1_1', value=str(options['OPTION_TEXT'].iloc[3]))
        options_1_2 = st.text_input("Welcher Kreis von betroffenen Personen?", key='option_1_2', value=str(options['OPTION_TEXT'].iloc[4]))
        options_1_3 = st.text_input("Welcher Zweck und welche konkrete Aufgabe der Datenverarbeitung besteht?", key='option_1_3', value=str(options['OPTION_TEXT'].iloc[5]))
    options_2 = st.container(border=True)
    options_2.header("Auftragsverhältnisse")
    with options_2.container(border=True):
        st.subheader("Unterauftrags-Verhältnisse")
        options_2_0 = st.toggle("Bestehen Unterauftrags-Verhältnisse?", key='option_2_0', value=eval(options['OPTION_TEXT'].iloc[6]))
    options_3 = st.container(border=True)
    options_3.header("Verarbeitung")
    with options_3.container(border=True):
        st.subheader("Intranet")
        options_3_0 = st.toggle("Sind die Daten oder Dienste über das Intranet erreichbar?", key='option_3_0', value=eval(options['OPTION_TEXT'].iloc[7]))
        options_3_1 = st.text_input("Daten", key='option_3_1', value=str(options['OPTION_TEXT'].iloc[8]))
        options_3_2 = st.text_input("Dienste", key='option_3_2', value=str(options['OPTION_TEXT'].iloc[9]))
    with options_3.container(border=True):
        st.subheader("Internet")
        options_3_3 = st.toggle("Sind die Daten oder Dienste über das Internet erreichbar?", key='option_3_3', value=eval(options['OPTION_TEXT'].iloc[10]))
        options_3_4 = st.text_input("Daten", key='option_3_4', value=str(options['OPTION_TEXT'].iloc[11]))
        options_3_5 = st.text_input("Dienste", key='option_3_5', value=str(options['OPTION_TEXT'].iloc[12]))
    with options_3.container(border=True):
        st.subheader("Anbieterzugang")
        options_3_6 = st.toggle("Muss beim Anbieter ein Zugang beantragt werden?", key='option_3_6', value=eval(options['OPTION_TEXT'].iloc[13]))
    with options_3.container(border=True):
        st.subheader("Webzugang")
        options_3_7 = st.toggle("Sind die Daten oder Dienste mit einem Webbrowser erreichbar?", key='option_3_7', value=eval(options['OPTION_TEXT'].iloc[14]))
        options_3_8 = st.text_input("Daten", key='option_3_8', value=str(options['OPTION_TEXT'].iloc[15]))
        options_3_9 = st.text_input("Dienste", key='option_3_9', value=str(options['OPTION_TEXT'].iloc[16]))
    with options_3.container(border=True):
        st.subheader("App-Zugang")
        options_3_10 = st.toggle("Sind die Daten oder Dienste mit einer App für Smartphones, Tablets oder PCs / Mac erreichbar?", key='option_3_10', value=eval(options['OPTION_TEXT'].iloc[17]))
        options_3_11 = st.toggle("Ist diese Lösung von BITMARK bereitgestellt?", key='option_3_11', value=eval(options['OPTION_TEXT'].iloc[18]))
        options_3_12 = st.toggle("Handelt es sich um einen Speicherplatz für Daten, der bereitgestellt wird?", key='option_3_12', value=eval(options['OPTION_TEXT'].iloc[19]))
    with options_3.container(border=True):
        st.subheader("KI-Tools und -Services")
        options_3_13 = st.toggle("Werden KI-Tools oder -Services verwendet?", key='option_3_13', value=eval(options['OPTION_TEXT'].iloc[20]))
    with options_3.container(border=True):
        st.subheader("Risiko- und Compliance-Bewertung")
        options_3_14 = st.toggle("Werden Services zur Risiko- und Compliance-Bewertung eingesetzt?", key='option_3_14', value=eval(options['OPTION_TEXT'].iloc[21]))
    with options_3.container(border=True):
        st.subheader("Herausgabe der Daten an US-Behörden")
        options_3_15 = st.toggle("Kann der Anbieter Daten unverschlüsselt an US-amrikanische Behörden übergeben?", key='option_3_15', value=eval(options['OPTION_TEXT'].iloc[22]))
    with options_3.container(border=True):
        st.subheader("Datenverarbeitung in der EU")
        options_3_16 = st.toggle(f"Werden alle {cloud}-SaaS-Dienste speziell für die Einhaltung von § 80 Abs. 2 SGB X konfiguriert, um eine Datenverarbeitung innerhalb der EU sicherzustellen?", key='option_3_16', value=eval(options['OPTION_TEXT'].iloc[23]))
    with options_3.container(border=True):
        st.subheader("Machine Learning")
        options_3_17 = st.toggle("Werden ML-Modelle zur Datenanalyse eingesetzt?", key='option_3_17', value=eval(options['OPTION_TEXT'].iloc[24]))
        options_3_18 = st.toggle("Wird sichergestellt, dass ML-Modelle in Infrastruktur, Tools und Workflows ausschließlich in der EU gehostet und genutzt werden?", key='option_3_18', value=eval(options['OPTION_TEXT'].iloc[25]))
        options_3_19 = st.toggle("Unterstützt der Anbieter die Erstellung von sicheren und konformen Machine Learning (ML)-Modellen im Sozial- und Gesundheitswesen?", key='option_3_19', value=eval(options['OPTION_TEXT'].iloc[26]))
    with options_3.container(border=True):
        st.subheader("Landing Zone")
        options_3_20 = st.toggle("Ist das Landing Zone Konzept Bestandteil der Cloud-Architektur?", key='option_3_20', value=eval(options['OPTION_TEXT'].iloc[27]))
        if cloud == 'AWS':
            options_3_21 = st.toggle("Folgt die Architektur dem AWS Well-Architected Framework?", key='option_3_21', value=eval(options['OPTION_TEXT'].iloc[28]))
        if cloud == 'Azure':
            options_3_22 = st.toggle("Folgt die Architektur dem Azure Well-Architected Framework?", key='option_3_22', value=eval(options['OPTION_TEXT'].iloc[29]))
        if cloud == 'Google Cloud':
            options_3_23 = st.toggle("Folgt die Architektur dem Google Cloud Architecture Framework?", key='option_3_23', value=eval(options['OPTION_TEXT'].iloc[30]))
        options_3_24 = st.text_input("Wie wurde die Landing Zone aufgebaut?", key='options_3_24', value=str(options['OPTION_TEXT'].iloc[31]))
        options_3_25 = st.toggle("Wurde dafür ein Tool genutzt?", key='options_3_25', value=eval(options['OPTION_TEXT'].iloc[32]))
        options_3_26 = st.text_input("Wie werden Identitäten und Zugriffsrechte innerhalb der Landing Zone verwaltet?", key='options_3_26', value=str(options['OPTION_TEXT'].iloc[33]))
        options_3_27 = st.toggle("Werden Maßnahmen zur Netzwerksicherheit in der Architektur der Landing Zone umgesetzt?", key='options_3_27', value=eval(options['OPTION_TEXT'].iloc[34]))
    with options_3.container(border=True):
        st.subheader("Rechtliche Anforderungen")
        options_3_28 = st.text_input(f"Welche Maßnahmen ergreift {cloud}, um sicherzustellen, dass die Datenverarbeitung und -speicherung den rechtlichen Anforderungen entspricht, auch im Hinblick auf US-amerikanische Auskunftsrechte?", key='option_3_28', value=str(options['OPTION_TEXT'].iloc[35]))
        options_3_29 = st.text_input(f"Wie adressiert {cloud} die Anforderungen an die Datenverarbeitung von SaaS-Diensten in Drittländern, insbesondere in Bezug auf das Schrems-II-Urteil und § 80 Abs. 2 SGB X?", key='options_3_29', value=str(options['OPTION_TEXT'].iloc[36]))
        options_3_30 = st.text_input(f"Wie können GKV-Träger sicherstellen, dass die Nutzung von {cloud}-SaaS-Diensten die Datenverarbeitung auf die EU beschränkt, in Übereinstimmung mit § 80 Abs. 2 SGB X?", key='options_3_30', value=str(options['OPTION_TEXT'].iloc[37]))
        options_3_31 = st.text_input("Wie wird die Übermittlung von Daten in Drittländer gehandhabt, insbesondere im Hinblick auf das Schrems-II-Urteil?", key='options_3_31', value=str(options['OPTION_TEXT'].iloc[38]))
        options_3_32 = st.text_input(f"Wie schützt {cloud} Kundendaten vor Zugriffen durch US-amerikanische Behörden?", key='options_3_32', value=str(options['OPTION_TEXT'].iloc[39]))
        options_3_33 = st.text_input(f"Wie unterstützt die {cloud}-Lösung die Einhaltung des Bundesdatenschutzgesetzes und der DSGVO?", key='options_3_33', value=str(options['OPTION_TEXT'].iloc[40]))
        options_3_34 = st.text_input(f"Wie gewährleistet {cloud} die Einhaltung der DSGVO und des BDSG für GKV-Daten?", key='options_3_34', value=str(options['OPTION_TEXT'].iloc[41]))
        options_3_35 = st.text_input(f"Wie gewährleistet {cloud} die Einhaltung von § 80 SGB X und DSGVO beim Umgang mit Sozialdaten?", key='options_3_35', value=str(options['OPTION_TEXT'].iloc[42]))
    with options_3.container(border=True):
        st.subheader("Compliance Anforderungen")
        options_3_36 = st.text_input(f"Wie können GKV-Träger {cloud}-Tools nutzen, um Compliance-Anforderungen zu überwachen und zu erfüllen?", key='options_3_36', value=str(options['OPTION_TEXT'].iloc[43]))
    with options_3.container(border=True):
        st.subheader("Backup-Strategien")
        options_3_37 = st.text_input(f"Welche {cloud}-Dienste nutzt die GKV im Bereich Hochverfügbarkeit?", key='options_3_37', value=str(options['OPTION_TEXT'].iloc[44]))
        options_3_38 = st.text_input(f"Welche {cloud}-Dienste nutzt die GKV im Bereich Disaster Recovery?", key='options_3_38', value=str(options['OPTION_TEXT'].iloc[45]))
    with options_3.container(border=True):
        st.subheader("Cloud Agnostik")
        options_3_39 = st.text_input("Welche Mechanismen bietet der Anbieter hinsichtlich eines Umzugs in eine andere Cloud (Cloud-Switching)?", key='options_3_39', value=str(options['OPTION_TEXT'].iloc[46]))
    with options_3.container(border=True):
        st.subheader("Faire Datennutzung")
        options_3_40 = st.toggle("Gibt es bereits Mechanismen zur Gewährleistung von European Data Act?", key='options_3_40', value=eval(options['OPTION_TEXT'].iloc[47]))
        options_3_41 = st.text_input(f"Wie nutzt die GKV {cloud}-Dienste, um eine feingranulare Zugriffssteuerung und Governance zu implementieren?", key='options_3_41', value=str(options['OPTION_TEXT'].iloc[48]))
    with options_3.container(border=True):
        st.subheader("IT-Sicherheit")
        st.write("Sicherheit")
        options_3_42 = st.text_input(f"Welche Maßnahmen ergreift {cloud} zum Schutz vor internen und externen Angriffen?", key='options_3_42', value=str(options['OPTION_TEXT'].iloc[49]))
        options_3_43 = st.text_input("Wie können GKV-Träger die Anforderungen an die physische Sicherheit und den Zugangsschutz in Rechenzentren überprüfen?", key='options_3_43', value=str(options['OPTION_TEXT'].iloc[50]))
        options_3_44 = st.text_input("Wie plant ihre Organisation, dies zu tun?", key='options_3_44', value=str(options['OPTION_TEXT'].iloc[51]))
        options_3_45 = st.toggle(f"Ist für die {cloud}-Lösung ein C5-Testat ausgestellt?", key='options_3_45', value=eval(options['OPTION_TEXT'].iloc[52]))
        st.write("Verschlüsselung")
        options_3_46 = st.text_input(f"Welche Maßnahmen trifft {cloud}, um Daten vor unbefugtem Zugriff durch Dritte, einschließlich {cloud} selbst, zu schützen?", key='options_3_46', value=str(options['OPTION_TEXT'].iloc[53]))
        options_3_47 = st.text_input(f"Wie unterstützt {cloud} die Verschlüsselung von Daten at-rest, in-transit und in-use?", key='options_3_47', value=str(options['OPTION_TEXT'].iloc[54]))
        options_3_48 = st.text_input(f"Wie unterstützt {cloud} die Verschlüsselung von Daten in-transit?", key='options_3_48', value=str(options['OPTION_TEXT'].iloc[55]))
        options_3_49 = st.text_input(f"Welche Bedeutung hat {cloud}-Verschlüsselungs-Dienst für die Verschlüsselung in-use?", key='options_3_49', value=str(options['OPTION_TEXT'].iloc[56]))
        options_3_50 = st.selectbox(label="Verschlüsselungs-Dienst", options=['AWS Nitro', 'Azure Managed HSM', 'Google Cloud KMS'], key='options_3_50', index=int(options['OPTION_TEXT'].iloc[57]))
        options_3_51 = st.text_input(f"Welche Mechanismen stellt die {cloud}-Lösung zur Verschlüsselung von Patientendaten zur Verfügung? Wie werden Schlüssel gemanaged?", key='options_3_51', value=str(options['OPTION_TEXT'].iloc[58]))

    st.header("Template")
    template = st.container(border=True)
    with template.container(border=True):
        st.subheader("Optionen")
        st.write("Bitte wähle die gewünschten Optionen aus 🔘")
        st.selectbox("Format", options=["Word"])
    with template.container(border=True):
        st.subheader("Absätze")
        st.write("Bitte wähle die Absätze aus 📒")
        paragraph_list = df["PARAGRAPH"].tolist()
        paragraph_title_list = df["PARAGRAPH_TITLE"].tolist()
        combined_list = [f"{p:<6} - {t}" for p, t in zip(paragraph_list, paragraph_title_list)]    
        chapters = st.multiselect("Absätze", options=combined_list, default=combined_list)
    submitted = st.form_submit_button("Template generieren")

if submitted:
    # Erase previous messages
    st.session_state.pop("langchain_messages", None)
    
    # Set up memory
    msgs = StreamlitChatMessageHistory(key="langchain_messages")
    if len(msgs.messages) == 0:
        msgs.add_ai_message(f"""Ich schreibe den Text in einer sachlichen und formellen
                                Form um und ersetze <Kunde> mit {kunde}, 
                                <Cloud-Anbieter> mit {cloud}.""")

    # Set up the LangChain, passing in Message History
    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", f"{system}"),
            MessagesPlaceholder(variable_name="history"),
            ("human", "{question}"),
        ]
    )

    # Setting the LLM
    if on:
        chain = prompt | ChatOpenAI(
            model="gpt-4o-mini",
            api_key=st.secrets["openai"]["key"]
        )
    else:
        server_url = f"{url}:{str(port)}/v1"
        chain = prompt | ChatOpenAI(
            base_url=server_url,
            model="llama-3-8b-chat-doctor-Q4_K_M_v2",
            temperature=0.5,
            max_tokens=4000,
            api_key="lm-studio"
      )

    chain_with_history = RunnableWithMessageHistory(
        chain,
        lambda session_id: msgs,
        input_messages_key="question",
        history_messages_key="history",
    )

    # Render current messages from StreamlitChatMessageHistory
    view_messages = st.status("Anzeige wird generiert...")
    with view_messages:
        for msg in msgs.messages:
            st.chat_message(msg.type).write(msg.content)

        # If user inputs a new prompt, generate and draw a new response
        for text in df["PARAGRAPH_TEXT"]:
            for chapter in chapters:
                if chapter[9:] in df["PARAGRAPH_TITLE"][df["PARAGRAPH_TEXT"] == text].to_string(index=False, header=False):
                    if '<Kundeninfo>' in text and web:
                        prompt = text.replace('<Kunde>', str(kunde)).replace('<Cloud-Anbieter>', str(cloud)).replace('<Kundeninfo>', str(kunde_info))
                    else:
                        prompt = text.replace('<Kunde>', str(kunde)).replace('<Cloud-Anbieter>', str(cloud))
                    if '<§' or '<Art.' in prompt:
                        for paragraph in paragraphs["PARAGRAPH"]:
                            # Checking for matching paragraph
                            if paragraph in prompt:
                                prompt = prompt.replace(f"<{paragraph}>", paragraphs[paragraphs['PARAGRAPH'] == paragraph].drop(columns=paragraphs.columns[-1]).to_string(index=False, header=False))
                                paragraph_info = web_scraper(paragraphs[paragraphs['PARAGRAPH'] == paragraph].drop(columns=paragraphs.columns[:2]).to_string(index=False, header=False))
                                paragraph_info = paragraph_info.replace('\n', ' ')
                                prompt += paragraph_info
                    if '<option_' in prompt:
                        for option in options['OPTION_DESC']: 
                            prompt = prompt.replace(f"<{option}>", str(options[options['OPTION_DESC'] == option].drop(columns=options.columns[:1]).to_string(index=False, header=False)))

                    st.chat_message("human").write(prompt)

                    # Note: new messages are saved to history automatically by Langchain during run
                    config = {"configurable": {"session_id": "any"}}
                    response = chain_with_history.invoke({"question": prompt}, config)
                    st.chat_message("ai").write(response.content)

    # Draw the messages at the end, so newly generated ones show up immediately
    view_chat_messages = st.expander("Zeige die Daten des Chatbots.")
    with view_chat_messages:
        """
        Message History initialized with:
        ```python
        msgs = StreamlitChatMessageHistory(key="langchain_messages")
        ```

        Contents of `st.session_state.langchain_messages`:
        """
        view_chat_messages.json(st.session_state.langchain_messages)

    # Convert to dataframe
    messages = st.session_state.langchain_messages
    anzeige_temp = pd.DataFrame(columns=['PARAGRAPH', 'PARAGRAPH_TITLE', 'PARAGRAPH_TEXT'])
    counter = -1
    paragraph = -1
    for index, message in enumerate(messages):
        for key, value in message:
            if key == "content":
                counter += 1
                if counter > 0 and counter % 2 == 0:
                    paragraph += 1
                    anzeige_temp = anzeige_temp._append(pd.DataFrame([{
                                                                        'PARAGRAPH': df['PARAGRAPH'][paragraph],
                                                                        'PARAGRAPH_TITLE': df['PARAGRAPH_TITLE'][paragraph],
                                                                        'PARAGRAPH_TEXT': value
                                                                      }]), 
                                                        ignore_index=True)

    st.dataframe(anzeige_temp)
    write_data(anzeige_temp, table_name='ANZEIGE_TEMP', database='OPENAI_DATABASE', schema='PUBLIC')
    with st.expander("Datenbankinhalt", expanded=False):
        df = load_data('OPENAI_DATABASE.PUBLIC.ANZEIGE_TEMP')
        st.dataframe(df)
    
    # Export to Word
    export_doc(anzeige_temp)
